from sys import stdin as s
from docx import Document
document = Document()
document.add_heading('Приглашение на мероприятие', 0)
s = s.read().split('\n')
p = document.add_paragraph('')
p.add_run('{0}, '.format(', '.join(s[2::]))).bold = True
p.add_run('приглашаем вас на меропрятие, которое состоится')
p.add_run('{0} {1}.\n'.format(s[0].lower(), s[1].lower())).italic = True
p.add_run('Спасибо за внимание!').gautami = True
document.save('test.docx')