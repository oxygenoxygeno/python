from docx import Document


def markdown_to_docx(text):
    document = Document()
    lines = text.split('\n')
    document.add_heading(lines[0], 0)
    for line in lines[1:]:
        if line:
            if line[:7].count('#') == 1:
                document.add_heading(line[2:], level=1)
            elif line[:7].count('#') == 2:
                document.add_heading(line[3:], level=2)
            elif line[:7].count('#') == 3:
                document.add_heading(line[4:], level=3)
            elif line[:7].count('#') == 4:
                document.add_heading(line[5:], level=4)
            elif line[:7].count('#') == 5:
                document.add_heading(line[6:], level=5)
            elif line[:7].count('#') == 6:
                document.add_heading(line[7:], level=6)
            elif str(line[:2]) == '- ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '* ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '+ ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and line[1] == '.':
                document.add_paragraph(line[3:], style='List Number')
            elif line[:3].count('_') == 1 or line[:3].count('*') == 1:
                document.add_paragraph().add_run(line[1:-1]).italic = True
            elif line[:3].count('_') == 2 or line[:3].count('*') == 2:
                document.add_paragraph().add_run(line[2:-2]).bold = True
            elif line[:3].count('_') == 3 or line[:3].count('*') == 3:
                runner = document.add_paragraph().add_run(line[3:-3])
                runner.bold = True
                runner.italic = True
            else:
                document.add_paragraph(line)
        else:
            document.add_paragraph()
    document.save('res.docx')